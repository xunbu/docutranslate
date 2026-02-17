# SPDX-FileCopyrightText: 2025 QinHan
# SPDX-License-Identifier: MPL-2.0
import base64
import io
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path

import markdown
from docutranslate.exporter.base import Exporter, ExporterConfig
from docutranslate.exporter.md.types import MD2DocxEngineType
from docutranslate.ir.document import Document
from docutranslate.ir.markdown_document import MarkdownDocument
from docutranslate.logger import global_logger


@dataclass(kw_only=True)
class MD2DocxExporterConfig(ExporterConfig):
    engine: MD2DocxEngineType = "auto"


def is_pandoc_available() -> bool:
    """检测pandoc是否可用"""
    return shutil.which("pandoc") is not None


def _md_to_docx_via_pandoc(md_content: str, logger=global_logger) -> bytes:
    """使用pandoc将markdown转换为docx"""
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        md_file = tmpdir / "input.md"
        docx_file = tmpdir / "output.docx"

        # 写入markdown文件（使用utf-8-sig确保中文兼容）
        md_file.write_text(md_content, encoding="utf-8-sig")

        # 执行pandoc转换
        try:
            result = subprocess.run(
                ["pandoc", str(md_file), "-o", str(docx_file)],
                capture_output=True,
                text=True,
                check=True
            )
            logger.info(f"Pandoc转换成功: {result.stdout}")
        except subprocess.CalledProcessError as e:
            logger.error(f"Pandoc转换失败: {e.stderr}")
            raise RuntimeError(f"Pandoc转换失败: {e.stderr}")

        return docx_file.read_bytes()


def _extract_base64_images(md_content: str):
    """提取markdown中的base64图片，返回(处理后的内容, 图片字典{索引: 图片信息})"""
    images = {}
    img_idx = 0

    # 匹配 ![alt](data:image/...) 格式，在原位置添加占位符
    def extract_md_image(match):
        nonlocal img_idx
        alt_text = match.group(1)
        data_url = match.group(2)
        # 提取base64数据
        header, b64data = data_url.split(',', 1)
        # 获取mime类型
        mime_type = header.split(';')[0].replace('data:', '')
        try:
            img_data = base64.b64decode(b64data)
            images[img_idx] = {
                'alt': alt_text,
                'mime': mime_type,
                'data': img_data
            }
            placeholder = f"<!--IMG_PLACEHOLDER_{img_idx}-->"
            img_idx += 1
            return placeholder
        except Exception:
            return ''  # 移除原文

    md_content = re.sub(r'!\[([^\]]*)\]\((data:image/[^)]+)\)', extract_md_image, md_content)

    return md_content, images


def _parse_markdown_table(table_text: str) -> list[list[str]]:
    """解析markdown表格"""
    # 检查是否是HTML表格
    if table_text.strip().startswith('<table>'):
        return _parse_html_table(table_text)

    lines = table_text.strip().split('\n')
    if len(lines) < 2:
        return []

    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            # 解析表格行
            cells = [cell.strip() for cell in line[1:-1].split('|')]
            rows.append(cells)

    return rows


def _parse_html_table(html_table: str) -> list[list[str]]:
    """解析HTML表格"""
    from bs4 import BeautifulSoup

    rows = []
    try:
        soup = BeautifulSoup(html_table, 'html.parser')
        for tr in soup.find_all('tr'):
            cells = []
            for td in tr.find_all(['td', 'th']):
                cells.append(td.get_text(strip=True))
            if cells:
                rows.append(cells)
    except Exception:
        pass

    return rows


def _add_image_to_docx(doc, img_info: dict, logger):
    """将图片添加到docx文档"""
    from docx.shared import Inches
    try:
        # 保存为临时文件
        ext = img_info['mime'].split('/')[-1]
        if ext == 'jpeg':
            ext = 'jpg'
        with tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False) as tmp:
            tmp.write(img_info['data'])
            tmp_path = tmp.name

        # 添加图片到docx
        if img_info['alt']:
            para = doc.add_paragraph()
            para.add_run(img_info['alt']).italic = True

        doc.add_picture(tmp_path, width=Inches(4))
        doc.add_paragraph()  # 图片后空行

        # 删除临时文件
        Path(tmp_path).unlink()
    except Exception as e:
        logger.warning(f"添加图片失败: {e}")


def _md_to_docx_via_python(md_content: str, logger=global_logger) -> bytes:
    """使用纯Python方式将markdown转换为docx"""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    # 提取并处理base64图片
    md_content, images = _extract_base64_images(md_content)

    # 移除HTML中的base64图片
    md_content = re.sub(r'<img[^>]*src="data:image/[^"]*"[^>]*>', '', md_content)

    # 将LaTeX公式转换为简单的文本表示（移除标记符号，保留公式内容）
    # 块公式 $$...$$ 或 \[...\]
    md_content = re.sub(r'\$\$([^$]+)\$\$', r'\1', md_content)
    md_content = re.sub(r'\\\[([^\\]+)\\\]', r'\1', md_content)
    # 行内公式 $...$ 或 \(...\)
    md_content = re.sub(r'\$([^$]+)\$', r'\1', md_content)
    md_content = re.sub(r'\\\(([^\\]+)\\\)', r'\1', md_content)

    # markdown转html
    extensions = [
        'markdown.extensions.tables',
        'markdown.extensions.fenced_code',
        'markdown.extensions.codehilite',
    ]
    html_content = markdown.markdown(md_content, extensions=extensions)

    # 创建docx文档
    doc = DocxDocument()

    # 处理表格 - 在解析markdown之前
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 检测图片占位符
        img_match = re.match(r'<!--IMG_PLACEHOLDER_(\d+)-->', line)
        if img_match:
            img_idx = int(img_match.group(1))
            if img_idx in images:
                _add_image_to_docx(doc, images[img_idx], logger)
            i += 1
            continue

        # 检测表格开始
        if line.startswith('|') and '|' in line[1:]:
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # 回退一行

            # 跳过分割线
            if len(table_lines) >= 2 and re.match(r'^\|[\s\-:|]+\|$', table_lines[1]):
                table_lines = table_lines[::2]  # 保留奇数行

            if len(table_lines) >= 1:
                table_data = _parse_markdown_table('\n'.join(table_lines))
                if table_data:
                    # 创建docx表格
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
                    table.style = 'Table Grid'

                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            if row_idx < len(table.rows) and col_idx < len(table.columns):
                                cell = table.cell(row_idx, col_idx)
                                cell.text = cell_text

                    doc.add_paragraph()  # 表格后空行
        else:
            # 非表格行，处理标题、列表等
            if not line:
                doc.add_paragraph()
            elif line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('- ') or line.startswith('* ') or line.startswith('+ '):
                doc.add_paragraph(line, style='List Bullet')
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '0.')):
                doc.add_paragraph(line, style='List Number')
            elif line.startswith('```') or line.startswith('~~~'):
                # 代码块
                code_lines = []
                i += 1
                while i < len(lines) and not (lines[i].strip().startswith('```') or lines[i].strip().startswith('~~~')):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    para = doc.add_paragraph('\n'.join(code_lines))
                    if para.runs:
                        para.runs[0].font.name = 'Courier New'
            else:
                # 检查行中是否有图片占位符
                if '<!--IMG_PLACEHOLDER_' in line:
                    # 分割行内容，处理占位符
                    parts = re.split(r'(<!--IMG_PLACEHOLDER_\d+-->)', line)
                    para = None
                    for part in parts:
                        img_match = re.match(r'<!--IMG_PLACEHOLDER_(\d+)-->', part)
                        if img_match:
                            img_idx = int(img_match.group(1))
                            if img_idx in images:
                                _add_image_to_docx(doc, images[img_idx], logger)
                        elif part.strip():
                            para = doc.add_paragraph(part)
                else:
                    # 处理普通段落
                    doc.add_paragraph(line)

        i += 1

    # 保存到bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


class MD2DocxExporter(Exporter):
    def __init__(self, config: MD2DocxExporterConfig | None = None):
        config = config or MD2DocxExporterConfig()
        super().__init__(config=config)
        self.engine = config.engine
        # 使用config中的logger，如果没有则使用全局logger
        self.logger = config.logger if hasattr(config, 'logger') and config.logger else global_logger

    def export(self, document: MarkdownDocument) -> Document:
        md_content = document.content.decode("utf-8")

        # 根据引擎选择转换方式
        engine = self.engine
        if engine == "auto":
            # 自动选择：有pandoc则用pandoc，否则用python
            engine = "pandoc" if is_pandoc_available() else "python"

        if engine == "pandoc":
            if not is_pandoc_available():
                self.logger.warning("Pandoc不可用，回退到纯Python模式")
                docx_bytes = _md_to_docx_via_python(md_content, self.logger)
            else:
                docx_bytes = _md_to_docx_via_pandoc(md_content, self.logger)
        else:
            docx_bytes = _md_to_docx_via_python(md_content, self.logger)

        # 转换为Document对象
        return Document.from_bytes(docx_bytes, suffix=".docx", stem=document.stem)
